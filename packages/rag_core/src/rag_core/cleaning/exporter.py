from __future__ import annotations

import re
from pathlib import Path

from rag_core.models import Document


class DualFormatExporter:
    def export(self, document: Document, output_dir: Path) -> tuple[Path, Path]:
        from docx import Document as WordDocument
        from docx.shared import Inches

        document_dir = output_dir / document.document_id
        document_dir.mkdir(parents=True, exist_ok=True)
        markdown_path = document_dir / f"{Path(document.filename).stem}.md"
        docx_path = document_dir / f"{Path(document.filename).stem}.docx"
        markdown_content = document.content
        for asset in document.assets:
            if asset.public_url:
                markdown_content = markdown_content.replace(asset.local_path, asset.public_url)
        markdown_path.write_text(markdown_content, encoding="utf-8")

        word = WordDocument()
        word.add_heading(document.title or Path(document.filename).stem, level=0)
        for line in document.content.splitlines():
            stripped = line.strip()
            heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            image = re.match(r"^!\[[^]]*]\(([^)]+)\)$", stripped)
            if heading:
                word.add_heading(heading.group(2), level=min(len(heading.group(1)), 6))
            elif image and Path(image.group(1)).exists():
                try:
                    word.add_picture(image.group(1), width=Inches(6.2))
                except Exception:
                    word.add_paragraph(stripped)
            elif stripped.startswith("| ") and stripped.endswith(" |"):
                word.add_paragraph(stripped)
            elif stripped:
                word.add_paragraph(stripped)
        word.save(docx_path)
        return markdown_path, docx_path
