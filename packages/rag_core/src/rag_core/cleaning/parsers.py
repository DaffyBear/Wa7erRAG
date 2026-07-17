from __future__ import annotations

import mimetypes
import uuid
from pathlib import Path

from rag_core.models import Asset, Document
from rag_core.utils import sha256_file, stable_id


class TextDocumentParser:
    extensions = {".txt", ".md", ".markdown"}

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.extensions

    def parse(self, path: Path) -> Document:
        content = path.read_text(encoding="utf-8-sig", errors="replace")
        return _new_document(path, content, title=_first_heading(content) or path.stem)


class HtmlDocumentParser:
    extensions = {".html", ".htm"}

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.extensions

    def parse(self, path: Path) -> Document:
        from bs4 import BeautifulSoup
        from markdownify import markdownify

        html = path.read_text(encoding="utf-8-sig", errors="replace")
        soup = BeautifulSoup(html, "html.parser")
        for element in soup(["script", "style", "noscript"]):
            element.decompose()
        title = soup.title.get_text(" ", strip=True) if soup.title else path.stem
        content = markdownify(str(soup.body or soup), heading_style="ATX")
        return _new_document(path, content, title=title)


class PdfDocumentParser:
    extensions = {".pdf"}

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.extensions

    def parse(self, path: Path) -> Document:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages: list[str] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(f"## Page {page_number}\n\n{text}")
        if not pages:
            raise ValueError(
                f"PDF contains no extractable text and requires OCR: {path.name}"
            )
        title = _pdf_title(reader) or path.stem
        document = _new_document(path, "\n\n".join(pages), title=title)
        document.metadata["page_count"] = len(reader.pages)
        return document


class DocxDocumentParser:
    extensions = {".docx"}

    def __init__(self, asset_root: Path) -> None:
        self.asset_root = asset_root

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.extensions

    def parse(self, path: Path) -> Document:
        from docx import Document as WordDocument
        from docx.document import Document as WordDocumentType
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        word = WordDocument(path)
        document_id = stable_id(str(path.resolve()), sha256_file(path))
        asset_dir = self.asset_root / document_id
        asset_dir.mkdir(parents=True, exist_ok=True)
        assets = self._extract_images(word, asset_dir)
        image_iter = iter(assets)
        lines: list[str] = []
        title = path.stem

        for block in _iter_blocks(word, Paragraph, Table, WordDocumentType):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                style_name = block.style.name.lower() if block.style else ""
                if text and style_name.startswith("heading"):
                    level = _heading_level(style_name)
                    lines.append(f"{'#' * level} {text}")
                    if title == path.stem:
                        title = text
                elif text:
                    lines.append(text)
                if "graphic" in block._p.xml or "drawing" in block._p.xml:
                    asset = next(image_iter, None)
                    if asset:
                        lines.append(f"![{Path(asset.local_path).stem}]({asset.local_path})")
            else:
                for row in block.rows:
                    values = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                    lines.append("| " + " | ".join(values) + " |")

        return Document(
            document_id=document_id,
            filename=path.name,
            source_path=path,
            content="\n\n".join(lines),
            title=title,
            checksum=sha256_file(path),
            assets=assets,
            metadata={"source_type": "docx"},
        )

    def _extract_images(self, word: object, asset_dir: Path) -> list[Asset]:
        assets: list[Asset] = []
        relationships = getattr(word.part, "rels", {})
        for relationship in relationships.values():
            if "image" not in relationship.reltype:
                continue
            blob = relationship.target_part.blob
            suffix = Path(relationship.target_ref).suffix or ".bin"
            target = asset_dir / f"{uuid.uuid4().hex}{suffix}"
            target.write_bytes(blob)
            content_type = mimetypes.guess_type(target.name)[0] or "application/octet-stream"
            assets.append(Asset(local_path=target.as_posix(), content_type=content_type))
        return assets


class ParserRegistry:
    def __init__(self, parsers: list[object]) -> None:
        self.parsers = parsers

    def parse(self, path: Path) -> Document:
        for parser in self.parsers:
            if parser.supports(path):
                return parser.parse(path)
        raise ValueError(f"Unsupported document type: {path.suffix}")

    def supports(self, path: Path) -> bool:
        return any(parser.supports(path) for parser in self.parsers)


def default_parser_registry(asset_root: Path) -> ParserRegistry:
    return ParserRegistry(
        [
            DocxDocumentParser(asset_root),
            PdfDocumentParser(),
            HtmlDocumentParser(),
            TextDocumentParser(),
        ]
    )


def _new_document(path: Path, content: str, title: str) -> Document:
    checksum = sha256_file(path)
    return Document(
        document_id=stable_id(str(path.resolve()), checksum),
        filename=path.name,
        source_path=path,
        content=content,
        title=title,
        checksum=checksum,
        metadata={"source_type": path.suffix.lower().lstrip(".")},
    )


def _first_heading(content: str) -> str:
    for line in content.splitlines():
        if line.lstrip().startswith("#"):
            return line.lstrip("# ").strip()
    return ""


def _heading_level(style_name: str) -> int:
    digits = "".join(character for character in style_name if character.isdigit())
    return max(1, min(6, int(digits or "1")))


def _iter_blocks(parent: object, paragraph_type: type, table_type: type, document_type: type):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    parent_element = parent.element.body if isinstance(parent, document_type) else parent._tc
    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield paragraph_type(child, parent)
        elif isinstance(child, CT_Tbl):
            yield table_type(child, parent)



def _pdf_title(reader: object) -> str:
    metadata = getattr(reader, "metadata", None)
    title = getattr(metadata, "title", "") if metadata else ""
    return str(title).strip() if title else ""
